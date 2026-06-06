#!/usr/bin/env python3
"""
Guía de Usuario — Quiniela Grupo MB Mundial 2026
Genera un DOCX profesional con tema mundialista y logo Grupo MB.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Paleta Mundialista ──
C_PRIMARY = RGBColor(0x0B, 0x3D, 0x0B)       # Verde pasto oscuro
C_ACCENT  = RGBColor(0xD4, 0xAF, 0x37)       # Dorado
C_LIGHT   = RGBColor(0xE8, 0xF5, 0xE9)       # Verde pasto claro
C_DARK    = RGBColor(0x1A, 0x1A, 0x1A)       # Casi negro
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_RED     = RGBColor(0xC0, 0x39, 0x2B)         # Rojo warning
C_GRAY    = RGBColor(0x66, 0x66, 0x66)

LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'GMB', 'RespaldosGMB', 'Nuevo logo MB.png')
OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'Guia-Quiniela-Grupo-MB-Mundial-2026.docx')

# ── Helpers ──

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_header(doc, logo_path):
    """Header con logo + título + línea de color"""
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)

    # Logo
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        run = cell.paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.2))

    # Title
    cell = table.cell(0, 1)
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("QUINIELA GRUPO MB")
    run.font.size = Pt(22)
    run.font.color.rgb = C_PRIMARY
    run.font.bold = True
    run = p.add_run("\nMUNDIAL FIFA 2026™")
    run.font.size = Pt(14)
    run.font.color.rgb = C_ACCENT
    run.font.bold = True

    # Línea separadora
    sep = doc.add_table(rows=1, cols=1)
    remove_table_borders(sep)
    set_cell_shading(sep.cell(0, 0), "0B3D0B")
    sep.cell(0, 0).height = Cm(0.15)
    sep.rows[0].height = Cm(0.15)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.color.rgb = C_PRIMARY
    run.font.bold = True

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = C_ACCENT
    run.font.bold = True

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.color.rgb = C_DARK
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = C_DARK

def add_numbered(doc, text, idx):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    r = p.add_run(f"  {idx}. ")
    r.font.bold = True
    r.font.color.rgb = C_ACCENT
    r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = C_DARK

def add_bullet(doc, text, icon="•"):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    r = p.add_run(f"  {icon} ")
    r.font.color.rgb = C_ACCENT
    r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = C_DARK

def add_note(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    r = p.add_run("  ⚠️ ")
    r.font.color.rgb = C_RED
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.color.rgb = C_GRAY
    r.font.size = Pt(10)
    r.font.italic = True

def add_spacer(doc, cm=0.5):
    p = doc.add_paragraph()
    p.space_before = Pt(cm * 14)
    p.space_after = Pt(0)

def create_data_table(doc, headers, rows, col_widths=None):
    """Tabla con header verde oscuro + filas alternadas"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "0B3D0B")
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.color.rgb = C_WHITE
        r.font.bold = True
        r.font.size = Pt(10)

    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            if ri % 2 == 0:
                set_cell_shading(cell, "E8F5E9")
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = C_DARK

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def add_footer(doc):
    sep = doc.add_table(rows=1, cols=1)
    remove_table_borders(sep)
    set_cell_shading(sep.cell(0, 0), "D4AF37")
    sep.cell(0, 0).height = Cm(0.1)
    sep.rows[0].height = Cm(0.1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4)
    r = p.add_run("Grupo MB — Quiniela Mundial 2026 — ¡Que gane el mejor! ⚽🏆")
    r.font.size = Pt(8)
    r.font.color.rgb = C_GRAY
    r.font.italic = True

# ── Document ──

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Header ──
add_header(doc, LOGO_PATH)

add_spacer(doc, 0.8)

# ── Bienvenida ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(6)
r = p.add_run("¡Bienvenido a la Quiniela del Mundial!")
r.font.size = Pt(16)
r.font.color.rgb = C_PRIMARY
r.font.bold = True

add_body(doc, "Participa prediciendo los resultados de los 104 partidos del Mundial FIFA 2026™. Acumula puntos y compite por premios increíbles. Es fácil, divertido y gratis.")

add_spacer(doc, 0.5)

# ── 1. Cómo Registrarte ──
add_section_heading(doc, "1. ¿Cómo Registrarte?")

add_numbered(doc, "Ve a la página de la Quiniela y haz clic en la pestaña de Registrarse.", 1)
add_numbered(doc, "Ingresa tu Nombre y tu Email.", 2)
add_numbered(doc, "Recibirás un código PIN de 6 dígitos. ¡GUÁRDALO!", 3)
add_numbered(doc, "En tu siguiente visita, usa tu Email y ese PIN para iniciar sesión.", 4)

add_note(doc, "Si pierdes tu PIN, contacta al administrador para generarte uno nuevo.")

add_spacer(doc, 0.3)

# ── 2. Cómo Hacer tus Predicciones ──
add_section_heading(doc, "2. Cómo Hacer tus Predicciones")

add_body(doc, "Una vez dentro de tu cuenta, ve a la sección ")
r = doc.paragraphs[-1].add_run('"Mis Predicciones"')
r.font.bold = True
r.font.color.rgb = C_PRIMARY
doc.paragraphs[-1].add_run(" y sigue estos pasos:")

add_numbered(doc, "Selecciona el resultado: Gana Local 🏠, Empate 🤝 o Gana Visita ✈️.", 1)
add_numbered(doc, "Opcionalmente, ingresa el marcador exacto (ej: 2-1) para ganar puntos extra.", 2)
add_numbered(doc, 'Haz clic en "💾 Guardar Predicciones" cuando termines.', 3)

add_note(doc, "Puedes editar tus predicciones cuantas veces quieras HASTA QUE INICIE el partido. Después de eso, se bloquean automáticamente.")

add_spacer(doc, 0.3)

# ── 3. Sistema de Puntos ──
add_section_heading(doc, "3. Sistema de Puntos (Escalonados)")

add_body(doc, "Los puntos aumentan conforme avanza el torneo. Atinar el resultado te da puntos base; atinar el marcador exacto te da puntos extra:")

create_data_table(doc,
    ["Fase", "Resultado Correcto", "Extra por Marcador"],
    [
        ["Fase de Grupos",      "3 pts", "+2 pts"],
        ["Ronda de 32",         "4 pts", "+3 pts"],
        ["Octavos de Final",    "5 pts", "+5 pts"],
        ["Cuartos de Final",    "7 pts", "+7 pts"],
        ["Semifinal",           "10 pts", "+10 pts"],
        ["Final",               "15 pts", "+15 pts"],
    ],
    col_widths=[2.5, 1.8, 1.8]
)

add_spacer(doc, 0.3)

add_subsection_heading(doc, "Ejemplo:")
add_bullet(doc, "Predices que México gana 2-0 vs Corea del Sur (Fase de Grupos)")
add_bullet(doc, "Si México gana 2-0 → ¡Ganas 5 pts! (3 resultado + 2 extra)")
add_bullet(doc, "Si México gana 1-0 → Ganas 3 pts (resultado correcto, marcador incorrecto)")
add_bullet(doc, "Si hay empate → 0 pts")

add_spacer(doc, 0.3)

# ── 4. Criterios de Desempate ──
add_section_heading(doc, "4. Criterios de Desempate")

add_body(doc, "Si dos o más participantes empatan en puntos totales, el desempate se resuelve así:")

add_numbered(doc, "Mayor número de marcadores exactos atinados.", 1)
add_numbered(doc, "Mayor número de resultados correctos.", 2)
add_numbered(doc, "Antigüedad de registro (quien se registró primero gana).", 3)

add_spacer(doc, 0.3)

# ── 5. Tiempo de Resultado ──
add_section_heading(doc, "5. Tiempo de Resultado")

add_body(doc, "Todos los resultados se consideran en ")
r = doc.paragraphs[-1].add_run("tiempo oficial de 90 minutos")
r.font.bold = True
r.font.color.rgb = C_PRIMARY
doc.paragraphs[-1].add_run(". No cuentan tiempos extras ni tanda de penales.")

add_note(doc, "Si un partido se define por penales, para la quiniela el resultado oficial es EMPATE.")

add_spacer(doc, 0.3)

# ── 6. Premios ──
add_section_heading(doc, "6. 🏆 Premios")

create_data_table(doc,
    ["Posición", "Premio"],
    [
        ["1° — 3° Lugar",   "Playera oficial + Balón oficial + Taza conmemorativa"],
        ["4° — 5° Lugar",   "Balón oficial del Mundial 2026"],
        ["6° — 10° Lugar",  "Taza conmemorativa"],
    ],
    col_widths=[2.0, 4.2]
)

add_spacer(doc, 0.5)

# ── 7. Leaderboard ──
add_section_heading(doc, "7. 📊 Leaderboard")

add_body(doc, "En la sección ")
r = doc.paragraphs[-1].add_run('"Leaderboard"')
r.font.bold = True
r.font.color.rgb = C_PRIMARY
doc.paragraphs[-1].add_run(" puedes ver la tabla de posiciones en tiempo real. Filtra por fase o consulta el ranking general. Tu posición aparece resaltada.")

add_spacer(doc, 0.3)

# ── 8. Reglas Importantes ──
add_section_heading(doc, "8. Reglas Importantes")

add_bullet(doc, "Cada partido acepta UNA predicción por usuario.")
add_bullet(doc, "Debes predecir al menos el resultado (Local, Empate o Visita). El marcador es opcional pero da puntos extra.")
add_bullet(doc, "Las predicciones se bloquean automáticamente cuando inicia el partido.")
add_bullet(doc, "El administrador puede bloquear temporalmente todas las predicciones (durante partidos en vivo).")
add_bullet(doc, "Los puntos se calculan automáticamente al terminar cada partido.")
add_bullet(doc, "No es necesario estar conectado al momento del partido; el sistema puntúa automáticamente.")

add_spacer(doc, 0.5)

# ── Cierre ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(12)
r = p.add_run("⚽ ¡Que comience el Mundial y que gane el mejor! 🏆")
r.font.size = Pt(14)
r.font.color.rgb = C_PRIMARY
r.font.bold = True

add_spacer(doc, 0.5)

# Footer
add_footer(doc)

# ── Save ──
output_path = os.path.abspath(OUTPUT)
doc.save(output_path)
print(f"✅ Guía generada: {output_path}")
